import os
import re
import threading
import tkinter as tk
from tkinter import messagebox, scrolledtext, simpledialog, ttk
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# ---------- 全局配置 ----------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
UA_FILE = os.path.join(SCRIPT_DIR, 'UA.txt')

DEFAULT_UA = (
    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
    'AppleWebKit/537.36 (KHTML, like Gecko) '
    'Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0'
)

HEADERS = {
    'User-Agent': DEFAULT_UA,
    'Accept-Language': 'zh-CN,zh;q=0.9'
}

# ---------- 辅助函数 ----------
def get_desktop_path():
    try:
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        if os.path.exists(desktop):
            return desktop
    except Exception:
        pass
    return SCRIPT_DIR

def is_valid_ua(ua):
    return bool(ua) and len(ua.strip()) >= 20

def load_user_agent():
    try:
        if os.path.exists(UA_FILE):
            with open(UA_FILE, 'r', encoding='utf-8') as f:
                ua = f.read().strip()
                return ua if is_valid_ua(ua) else None
    except Exception:
        pass
    return None

def save_user_agent(ua):
    try:
        with open(UA_FILE, 'w', encoding='utf-8') as f:
            f.write(ua.strip())
        return True
    except Exception:
        return False

# ---------- 爬虫核心 ----------
def fetch_word_info(word):
    url = f'https://www.youdao.com/result?word={word}&lang=en'
    try:
        resp = requests.get(url, headers=HEADERS, timeout=10)
        resp.raise_for_status()
    except requests.RequestException as e:
        raise RuntimeError(f"网络请求失败: {e}")

    soup = BeautifulSoup(resp.text, 'html.parser')

    # 单词
    title_tag = soup.find('div', class_='title')
    if title_tag:
        direct_text = title_tag.find(string=True, recursive=False)
        actual_word = direct_text.strip() if direct_text else word
    else:
        actual_word = word

    # 释义
    pos_trans = []
    basic_ul = soup.select_one('.trans-container .basic')
    if basic_ul:
        for li in basic_ul.find_all('li', class_='word-exp'):
            pos_tag = li.find('span', class_='pos')
            trans_tag = li.find('span', class_='trans')
            if pos_tag and trans_tag:
                pos_trans.append((pos_tag.get_text(strip=True), trans_tag.get_text(strip=True)))

    # 短语（放在例句之前）
    phrases = []
    phrs_section = soup.find('div', class_='phrs')
    if phrs_section:
        for li in phrs_section.select('ul.trans-container li'):
            phr_a = li.find('a', class_='point')
            phr_trans = li.find('span', class_='phr_trans')
            if phr_a and phr_trans:
                phrases.append((phr_a.get_text(strip=True), phr_trans.get_text(strip=True)))

    # 例句（修复粘连）
    sentences = []
    sents_div = soup.find('div', class_='blng_sents_part')
    if sents_div:
        for li in sents_div.select('ul li.mcols-layout'):
            eng_tag = li.select_one('.sen-eng')
            chn_tag = li.select_one('.sen-ch')
            if eng_tag and chn_tag:
                for b in eng_tag.find_all('b'):
                    b.insert_before(' ')
                    b.insert_after(' ')
                eng = eng_tag.get_text()
                eng = re.sub(r'\s+', ' ', eng).strip()
                eng = re.sub(r'\s([.,;:!?])', r'\1', eng)
                chn = chn_tag.get_text(strip=True)
                sentences.append((eng, chn))

    return {
        'word': actual_word,
        'pos_trans': pos_trans,
        'phrases': phrases,
        'sentences': sentences
    }

def process_single_word(args):
    """并发任务函数：抓取一个单词并应用释义上限"""
    index, word, limit = args
    try:
        info = fetch_word_info(word)
        pos_trans = []
        for pos, trans in info['pos_trans']:
            if limit > 0:
                items = trans.split('；')
                trans = '；'.join(items[:limit])
            pos_trans.append((pos, trans))

        return (index, {
            'word': info['word'],
            'pos_trans': pos_trans,
            'phrases': info['phrases'],
            'sentences': info['sentences']
        })
    except Exception as e:
        return (index, {
            'word': word,
            'pos_trans': [],
            'phrases': [],
            'sentences': [],
            'error': str(e)
        })

# ---------- GUI ----------
class WordApp:
    def __init__(self, root):
        self.root = root
        root.title("气死雷颜器")
        root.geometry("500x450")
        root.resizable(False, False)

        # 图标
        icon_path = os.path.join(SCRIPT_DIR, 'favicon.ico')
        try:
            root.iconbitmap(icon_path)
        except Exception:
            pass

        # 单词输入
        tk.Label(root, text="输入单词（每行一个，或用逗号/空格分隔）：").pack(pady=(10, 0))
        self.word_text = scrolledtext.ScrolledText(root, height=10, width=60)
        self.word_text.pack(pady=5)

        # 释义上限
        limit_frame = tk.Frame(root)
        limit_frame.pack(pady=5)
        tk.Label(limit_frame, text="释义个数上限（0 为不限制）：").pack(side=tk.LEFT)
        self.limit_entry = tk.Entry(limit_frame, width=5)
        self.limit_entry.insert(0, "0")
        self.limit_entry.pack(side=tk.LEFT, padx=5)

        # 并行数设定
        workers_frame = tk.Frame(root)
        workers_frame.pack(pady=5)
        tk.Label(workers_frame, text="并行线程数（1~10，推荐5）：").pack(side=tk.LEFT)
        self.workers_entry = tk.Entry(workers_frame, width=5)
        self.workers_entry.insert(0, "5")
        self.workers_entry.pack(side=tk.LEFT, padx=5)

        # 进度条
        self.progress = ttk.Progressbar(root, mode='determinate', length=400)
        self.progress.pack(pady=10)

        # 开始按钮
        self.start_btn = tk.Button(root, text="开始生成 Word 文档", command=self.start)
        self.start_btn.pack(pady=10)

        # 状态栏
        self.status_label = tk.Label(root, text="就绪", fg="gray")
        self.status_label.pack()

        #信息栏
        self.info_label = tk.Label(root, text="本程序由杨镇源开发，仅供完成雷颜布置的20篇20词积累使用。\n请勿向雷颜透露、将本程序另作他用等，最终解释权归作者本人所有。\n© anghenuan", fg="gray")
        self.info_label.pack(pady=15)

    def parse_words(self):
        raw = self.word_text.get("1.0", tk.END)
        words = re.split(r'[,\s]+', raw)
        return [w.strip() for w in words if w.strip()]

    def start(self):
        # UA 检查
        ua = load_user_agent()
        while ua is None:
            user_input = simpledialog.askstring(
                "User‑Agent 配置",
                "未找到有效的 User‑Agent 文件。\n请输入完整的 User‑Agent 字符串（至少 20 个字符）：",
                parent=self.root
            )
            if user_input is None:
                messagebox.showwarning("操作取消", "未配置 User‑Agent，无法继续运行。")
                return
            user_input = user_input.strip()
            if not is_valid_ua(user_input):
                messagebox.showwarning("格式错误", "User‑Agent 长度不足或为空，请重新输入。")
                continue
            if save_user_agent(user_input):
                ua = user_input
                break
            else:
                messagebox.showerror("保存失败", "无法写入 UA.txt，请检查文件权限。")
                return

        global HEADERS
        HEADERS['User-Agent'] = ua

        words = self.parse_words()
        if not words:
            messagebox.showwarning("提示", "请至少输入一个单词")
            return
        try:
            limit = int(self.limit_entry.get().strip())
        except ValueError:
            limit = 0

        # 解析并行数
        try:
            max_workers = int(self.workers_entry.get().strip())
            if max_workers < 1:
                max_workers = 1
            elif max_workers > 10:
                max_workers = 10
        except ValueError:
            max_workers = 5  # 默认值

        self.max_workers = max_workers

        # 准备并发任务
        self.total_words = len(words)
        self.completed = 0
        self.progress['maximum'] = self.total_words
        self.progress['value'] = 0
        self.status_label.config(text="正在抓取... 0/%d" % self.total_words)
        self.start_btn.config(state=tk.DISABLED)

        # 构建参数列表
        tasks = [(i, w, limit) for i, w in enumerate(words, 1)]
        # 启动并发线程
        threading.Thread(target=self.run_concurrent, args=(tasks,), daemon=True).start()

    def run_concurrent(self, tasks):
        results_dict = {}
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {executor.submit(process_single_word, task): task[0] for task in tasks}
            for future in as_completed(futures):
                idx, result = future.result()
                results_dict[idx] = result
                self.root.after(0, self.update_progress)
        sorted_results = [results_dict[i] for i in sorted(results_dict.keys())]
        self.root.after(0, self.finish_up, sorted_results)

    def update_progress(self):
        self.completed += 1
        self.progress['value'] = self.completed
        self.status_label.config(text=f"正在抓取... {self.completed}/{self.total_words}")

    def finish_up(self, results):
        for i, entry in enumerate(results, 1):
            entry['index'] = i
        self.generate_word(results)
        self.status_label.config(text="生成完毕，文件已保存到桌面")
        self.start_btn.config(state=tk.NORMAL)
        messagebox.showinfo("完成", "Word 文档已生成到桌面：单词积累.docx")

    def generate_word(self, entries):
        doc = Document()
        doc.core_properties.author = '由气死雷颜器自动生成'   # 设置作者
        doc.core_properties.comments = ''   # 清空备注
        '''
        from datetime import datetime
        now = datetime.now()
        # 1. 获取系统当前本地时间，并格式化为不带时区的字符串（如 2026-07-05T10:30:00）
        local_now_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
        # 2. 获取核心属性元素
        core_element = doc.core_properties._element
        # 命名空间映射
        nsmap = {'dcterms': 'http://purl.org/dc/terms/'}
        # 3. 处理 created 元素
        created_el = core_element.find('dcterms:created', nsmap)
        if created_el is None:
            created_el = core_element.makeelement('{http://purl.org/dc/terms/}created', {})
            core_element.append(created_el)
        created_el.text = local_now_str
        # 4. 处理 modified 元素
        modified_el = core_element.find('dcterms:modified', nsmap)
        if modified_el is None:
            modified_el = core_element.makeelement('{http://purl.org/dc/terms/}modified', {})
            core_element.append(modified_el)
        modified_el.text = local_now_str
        '''
        style = doc.styles['Normal']
        font = style.font
        font.name = '等线'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        #style.paragraph_format.line_spacing = 1.5

        for entry in entries:
            # 首行：序号 + 单词
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            #p.paragraph_format.line_spacing = 1.5

            run_title = p.add_run(f"{entry.get('index', '?')}. {entry['word']}   ")
            run_title.bold = True
            self.set_font(run_title)

            if 'error' in entry:
                run_err = p.add_run(f"（抓取失败: {entry['error']}）")
                self.set_font(run_err)
                continue

            # 词性 + 释义
            for pos, trans in entry.get('pos_trans', []):
                run_pos = p.add_run(f" {pos} ")
                run_pos.bold = True
                self.set_font(run_pos)
                run_trans = p.add_run(f"{trans}  ")
                self.set_font(run_trans)

            # 短语（提前至例句之前）
            if entry.get('phrases'):
                for idx, (phr, trans) in enumerate(entry['phrases'], 1):
                    p_phr = doc.add_paragraph()
                    p_phr.paragraph_format.left_indent = Pt(20)
                    p_phr.paragraph_format.first_line_indent = Pt(-10)
                    #p_phr.paragraph_format.line_spacing = 1.3
                    run_label = p_phr.add_run(f"短语{idx}: ")
                    run_label.bold = True
                    self.set_font(run_label)
                    run_content = p_phr.add_run(f"{phr}  {trans}")
                    self.set_font(run_content)

            # 例句
            if entry.get('sentences'):
                for idx, (eng, chn) in enumerate(entry['sentences'], 1):
                    p_ex = doc.add_paragraph()
                    p_ex.paragraph_format.left_indent = Pt(20)
                    p_ex.paragraph_format.first_line_indent = Pt(-10)
                    #p_ex.paragraph_format.line_spacing = 1.3
                    run_label = p_ex.add_run(f"例{idx}: ")
                    run_label.bold = True
                    self.set_font(run_label)
                    run_content = p_ex.add_run(f"{eng}  {chn}")
                    self.set_font(run_content)
            else:
                p_ex = doc.add_paragraph()
                p_ex.paragraph_format.left_indent = Pt(20)
                run_label = p_ex.add_run("例: ")
                run_label.bold = True
                self.set_font(run_label)
                run_none = p_ex.add_run("（无例句）")
                self.set_font(run_none)

            # 分隔线
            '''
            if entry != entries[-1]:
                div = doc.add_paragraph('─' * 50)
                div.paragraph_format.space_before = Pt(4)
                div.paragraph_format.space_after = Pt(4)
                div.alignment = 1
                for run in div.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            '''

        output_path = os.path.join(get_desktop_path(), "单词积累.docx")
        doc.save(output_path)

    def set_font(self, run):
        run.font.name = '等线'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

if __name__ == "__main__":
    root = tk.Tk()
    app = WordApp(root)
    root.mainloop()
